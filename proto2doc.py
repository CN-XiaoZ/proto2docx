from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Inches
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt
from docx.oxml.ns import qn

# 解析一个<table>标签, 转成一个列表, 一个成员就是一个行. 
class TableParser:
    def __init__(self, table) -> None:
        self.table = table
        self.row_list = []
        self.thead = [td.text for td in table.thead.find_all('td')]
        for tr in table.tbody.find_all('tr'):
            self.row_list.append([x for x in tr.find_all('td')])

    def row_iter(self):
        for x in self.row_list:
            yield(x)

#将Service下面两个表, 每个method组成一个MethodInfo
class MethodInfo:
    def __init__(self, method_row, method_path) -> None:
        self.name = method_row[0].string
        self.http_method = method_path[1].string
        self.http_path = method_path[2].string
        self.req = method_row[1].a["href"].strip("#")
        self.resp = method_row[2].a["href"].strip("#")
        self.desc = method_row[3].string
    def __str__(self) -> str:
        return "%s %s REQ:%s RESP:%s [%s]"%(self.http_method,self.http_path, self.req, self.resp, self.desc)

class ServiceInfo:
    def __init__(self, parser, table_method, table_path) -> None:
        #将Service下所有method组成一个列表放在这里
        self.method_list = []
        for path in table_path.row_list:
            name = path[0].string
            for info in table_method.row_list:
                if name == info[0].string:
                    self.method_list.append(MethodInfo(info, path))

class DocCreater:
    table_desc = {}
    buildin_types = ["double", "float", "int32", "int64", "uint32", "uint64",
                     "sint32", "sint64", "fixed32", "fixed64", "sfixed32", "sfixed64",
                     "bool","string","bytes"]
    def __init__(self, doc_name, parser) -> None:
        # 打开文档，写一下头之类的，
        self.doc_name = doc_name
        self.table_desc = [["参数", "类型", "长度", "是否必须", "描述"], self.write_table]
        self.doc = Document()
        self.parser = parser
        self.doc.add_heading('OpenApi 单机版', 0)
    
    def is_link(self, item):
        if item.a["href"].strip("#") in self.buildin_types:
            return False
        return True

    def write_table(self, doc_table, table_name, method_type, prefix = ""):
        table = TableParser(self.parser.get_table_by_name(table_name))
        if method_type == "resp":
            if table.row_list[0][0].string != "data" and prefix == "":
                self.write_line_simple(doc_table, ["data", "Object", "-", "是", "数据"])
                prefix += "data"
        for row in table.row_list:
            self.write_line(doc_table, row, method_type, prefix)

    def write_line_simple(self, doc_table, vals):
        cells = doc_table.add_row().cells
        for i,val in enumerate(vals):
            cells[i].text = val

    def write_line(self, table, row, method_type, prefix = ""):
        next = None
        gotag = None
        tag = ""
        cells = table.add_row().cells
        if prefix != "":
            prefix += "."
        cells[0].text = prefix + row[0].string
        if str(row[2].string).strip() == "repeated":
            tag = "[]"
        link = row[1].a["href"].strip("#")
        if link not in self.buildin_types:
            next = link
            cells[1].text = "Object Array"
        else:
            cells[1].text = link
        cells[2].text = "-"
        if cells[0].text.strip() == "data" and method_type == "resp":
            cells[3].text = "是"
        else:
            cells[3].text = "否"
        desc = row[3].p.string

        if "@gotags" in desc:
            cells[4].text = desc.split("@")[0].strip()
            gotag = desc.split("@")[1].strip()
            if "required" in gotag:
                cells[3].text = "是"
        else:
            cells[4].text = desc

        if next:
            self.write_table(table, next, method_type, prefix = cells[0].text + tag)


    def write_method_table(self, name, method_type, add = []):
        table_header = self.table_desc[0]
        write_func = self.table_desc[1]

        doc_table = self.doc.add_table(rows=1, cols=len(table_header))
        for i, cell in enumerate(doc_table.rows[0].cells):
            cell.text = table_header[i]
        if add != []:
            for line in add:
                self.write_line_simple(doc_table, line)
        write_func(doc_table, name, method_type)

    # 向文档里写入一个method.
    def write_method(self, method):
        self.doc.add_heading(method.name, level=1)
        self.doc.add_heading("1. 接口说明", level=2)
        self.doc.add_paragraph(method.desc)
        self.doc.add_heading("2. 访问说明", level=2)
        self.doc.add_paragraph("Path: " + method.http_path + "\nMethod: " + method.http_method + "\nContent-type: application/json")
        self.doc.add_heading("3. 请求参数", level=2)
        self.write_method_table(method.req, "req")
        self.doc.add_heading("4. 响应参数", level=2)
        self.write_method_table(method.resp,"resp", add = [["code", "int","-","是","0: 成功, 其余详见错误码说明"],["message", "string","-","是","应答描述"]])
        self.doc.add_heading("5. 示例代码", level=2)
        self.doc.add_paragraph("")
        self.doc.add_heading("6. 请求报文", level=2)
        self.doc.add_paragraph("")
        self.doc.add_heading("7. 响应报文", level=2)
        self.doc.add_paragraph("")
        
        # self.doc.add_page_break()

    def save(self):
        for style in self.doc.styles:
            if style.type == WD_STYLE_TYPE.PARAGRAPH and style.name.startswith('Heading'):
                style.font.name = '微软雅黑'  # 将标题的中文字体改为宋体

        # 修改英文字体
        for paragraph in self.doc.paragraphs:
            for run in paragraph.runs:
                if run.font.name != '微软雅黑':
                    run.font.name = 'Courier New'  # 将英文的字体改为Courier New
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')  # 将中文字体改为宋体

        # 修改表格样式
        for table in self.doc.tables:
            # 将表格整体居中
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

            # 将表格中所有单元格的默认对齐方式改为左对齐，并将字体改为宋体
            for row in table.rows:
                for cell in row.cells:
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                    for run in cell.paragraphs[0].runs:
                        run.font.name = 'Courier New'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

            # 将表格首行对齐方式改为居中，并将字体改为宋体和加粗
            for cell in table.rows[0].cells:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                cell.paragraphs[0].runs[0].font.bold = True
                cell.paragraphs[0].runs[0].font.name = '微软雅黑'
                cell.paragraphs[0].runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

        self.doc.save(self.doc_name)

class HtmlParser:
    def __init__(self, html= "index.html") -> None:
        self.soup = BeautifulSoup(open(html),features="html.parser")
        self.content_list = []
        self.service_list = []
        self.parse_toc()
        self.process_to_doc()

    def process_to_doc(self):
        for service in self.service_list_iter():
            service_table = self.get_table_by_name(service["link"])
            service_path_table = service_table.find_next_sibling('table')
            self.service_list.append(ServiceInfo(self, TableParser(service_table), TableParser(service_path_table)))
            # 这个break是test， 只运行一次， 正式要去掉
            break

    def parse_toc(self):
        toc_ul = self.soup.find('ul', {'id': 'toc'})

        for li in toc_ul.find_all('li'):
            link_type = li.a.span
            if link_type != None:
                # {'link': 'api.license.v1.LicenseService', 'name': 'LicenseService', 'type': 'S'}
                self.content_list.append({
                    "link": li.a["href"].strip("#"),
                    "name":li.a.contents[1],
                    "type":link_type.string
                })

    def service_list_iter(self):
        for link in self.content_list:
            if link["type"] == "S":
                yield link

    def get_table_by_name(self, name):
        return self.soup.find('h3', id=name).find_next_sibling('table')

if __name__ == "__main__":
    hp = HtmlParser()
    dc = DocCreater("output.doc", hp)
    for service in hp.service_list:
        for method in service.method_list:
            dc.write_method(method)
    dc.save()